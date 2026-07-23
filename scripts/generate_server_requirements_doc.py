from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "SERVER_XARAKTERISTIKASI_500_ONLINE_BIG_DATA.docx"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_fill(cell, "123052")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(label))
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        if widths:
            cell.width = widths[idx]

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_fill(cell, "F3F6F9")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8.5)
            if widths:
                cell.width = widths[idx]
    document.add_paragraph()
    return table


def add_bullet(document, text, bold_prefix=None):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        run = paragraph.add_run(first + ":")
        run.bold = True
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sahifa ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


document = Document()
section = document.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.55)
section.right_margin = Cm(1.55)
section.header_distance = Cm(0.7)
section.footer_distance = Cm(0.7)

styles = document.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in (
    ("Title", 24, "123052"),
    ("Heading 1", 16, "123052"),
    ("Heading 2", 12, "117A65"),
):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

header = section.header.paragraphs[0]
header.text = "STATISTIKA QO'MITASI  |  DATA WAREHOUSE INFRASTRUKTURASI"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header.runs:
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("607080")
add_page_number(section.footer.paragraphs[0])

title = document.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("500 ta online foydalanuvchi va katta hajmdagi data uchun\nserver va servis talablari")

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Enterprise Data Warehouse va Data Lake production infratuzilma tavsiyasi").bold = True
subtitle.add_run("\nVersiya 2.0 | 20.07.2026")

notice = document.add_table(rows=1, cols=1)
notice.alignment = WD_TABLE_ALIGNMENT.CENTER
notice.style = "Table Grid"
cell = notice.cell(0, 0)
set_cell_fill(cell, "EAF4F1")
set_cell_margins(cell, 150, 180, 150, 180)
paragraph = cell.paragraphs[0]
run = paragraph.add_run(
    "Profil: o'rtacha 500 ta online user, 800 ta peak user, 100-250 GB/kun boshlang'ich ingest, "
    "ikki yillik Data Lake retention va katta hajmdagi parallel ETL. "
    "0.5-2 TB/kun yuklama uchun hujjatdagi PB-scale profil qo'llanadi."
)
run.bold = True
run.font.color.rgb = RGBColor.from_string("145A4A")
document.add_paragraph()

document.add_heading("1. Tizim yuklamasi va asosiy farazlar", level=1)
document.add_paragraph(
    "Bu sizing faqat web foydalanuvchilar soniga emas, ma'lumot hajmi va qayta ishlash murakkabligiga asoslangan. "
    "Frontend animatsiyasi client brauzerida ishlaydi. Asosiy infratuzilma yuki Kafka ingestion, MinIO object write, "
    "Spark transform, Trino ad-hoc query, ClickHouse dashboard query va eksport jarayonlaridan keladi."
)
add_table(
    document,
    ["Ko'rsatkich", "Boshlang'ich katta-data target"],
    [
        ("Online foydalanuvchi", "o'rtacha 500; peak 800; stress 1000"),
        ("API yuklamasi", "75-150 request/second"),
        ("Kunlik yangi data", "100-250 GB/kun; peak 500 GB/kun"),
        ("Taxminiy record", "100-500 million row/kun; row widthga bog'liq"),
        ("Data Lake retention", "Raw 2 yil; Landing 7-30 kun"),
        ("ClickHouse hot retention", "6-12 oy"),
        ("Parallel analytic query", "50-100 running; qolganlari resource queue"),
        ("Parallel Spark pipeline", "8-20 job; qolganlari priority queue"),
        ("Cached dashboard", "p95 < 3 soniya"),
        ("Oddiy API", "p95 < 1 soniya"),
        ("Mavjudlik", "kamida 99.9%; kritik qatlamlar HA"),
    ],
    [Inches(2.8), Inches(3.8)],
)

document.add_heading("2. Storage capacity qanday hisoblanadi", level=1)
document.add_paragraph(
    "Foydalanuvchi soni storage hajmini belgilamaydi. Asosiy formula:"
)
formula = document.add_paragraph()
formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
formula_run = formula.add_run(
    "Raw capacity = kunlik ingest x retention kuni x zone koeffitsiyenti / EC samaradorligi / 70% fill limit"
)
formula_run.bold = True
formula_run.font.color.rgb = RGBColor.from_string("123052")
document.add_paragraph(
    "Misol: 250 GB/kun x 730 kun x 1.8 zone/version koeffitsiyenti / 0.75 erasure-code samaradorligi / 0.70 fill limit "
    "= taxminan 625 TB raw disk. Shu sababli boshlang'ich MinIO pool 768 TB raw qilib olinadi. "
    "Bu raqamlar real Parquet siqilishi, fayl versiyalari va retention siyosati bilan load-testdan keyin aniqlashtiriladi."
)
add_table(
    document,
    ["Saqlash qatlami", "Retention", "Format", "Maqsad"],
    [
        ("Landing", "7-30 kun", "Original JSON/CSV/Excel", "Manbadan kelgan o'zgarmagan payload"),
        ("Raw", "2 yil", "Parquet + compression", "Qayta ishlash mumkin bo'lgan asosiy tarix"),
        ("Prepared", "90-365 kun", "Apache Iceberg/Parquet", "Imputatsiya, versiya va audit snapshot"),
        ("Curated", "1-2 yil", "Apache Iceberg/Parquet", "Conformed model va Trino query"),
        ("ClickHouse Hot DWH", "6-12 oy", "MergeTree", "Dashboard, KPI va tezkor analytics"),
        ("Backup/Archive", "siyosat bo'yicha", "Object backup", "DR va uzoq muddatli tiklash"),
    ],
    [Inches(1.55), Inches(1.15), Inches(1.7), Inches(2.5)],
)

document.add_heading("3. Ikki xil katta-data profili", level=1)
add_table(
    document,
    ["Parametr", "Large-A: tavsiya etilgan start", "Large-B: PB-scale"],
    [
        ("Kunlik ingest", "100-250 GB", "0.5-2 TB"),
        ("Raw retention", "2 yil", "2-5 yil"),
        ("MinIO raw capacity", "768 TB raw; ~400 TB safe usable", "1-4 PB safe usable"),
        ("ClickHouse", "6 node: 3 shard x 2 replica", "8-16 node; shardlar kengayadi"),
        ("Spark/Trino compute", "8 ta katta compute node", "16-32 compute node"),
        ("Kubernetes worker", "18-22 node; storage tashqarida", "28-45 node; storage tashqarida"),
        ("Umumiy CPU", "600-900 vCPU", "1200-2500 vCPU"),
        ("Umumiy RAM", "2.5-4 TB", "6-12 TB"),
        ("Tarmoq", "25 Gbit/s; storage 40 Gbit/s", "40-100 Gbit/s"),
    ],
    [Inches(1.6), Inches(2.6), Inches(2.4)],
)
document.add_paragraph(
    "Loyiha uchun Large-A profildan boshlash tavsiya etiladi. Agar real ingest 250 GB/kun yoki MinIO 60% fill darajasidan oshsa, "
    "oldindan rejalangan yangi server pool bilan Large-B profiliga kengaytiriladi."
)

document.add_heading("4. Large-A uchun servislar bo'yicha sizing", level=1)
service_rows = [
    ("NGINX / HAProxy", "2", "4 vCPU", "8 GB", "TLS, WAF, routing, rate limit"),
    ("Next.js frontend", "4", "4 vCPU", "8 GB", "Stateless; HPA 4-12 pod"),
    ("FastAPI Gateway", "6", "8 vCPU", "16 GB", "Job qabul qiladi; katta payloadni olib o'tmaydi"),
    ("Job status worker", "8-12", "8 vCPU", "16 GB", "Metadata, retry, export va queue"),
    ("Redis Cluster", "3-6", "8 vCPU", "32 GB", "Cache, session, queue va distributed lock"),
    ("Keycloak", "3", "8 vCPU", "16 GB", "SSO, MFA, IAM va RBAC"),
    ("PostgreSQL HA", "3", "16 vCPU", "64 GB", "2 TB NVMe; primary + 2 replica"),
    ("PgBouncer", "2", "4 vCPU", "8 GB", "Connection pooling"),
    ("Kafka Broker", "5", "16 vCPU", "64 GB", "4 TB NVMe; RF=3, min ISR=2"),
    ("Apache NiFi", "3", "16 vCPU", "64 GB", "2 TB NVMe; back-pressure va provenance"),
    ("Airflow Scheduler", "2", "8 vCPU", "16 GB", "HA scheduler va DAG processor"),
    ("Airflow Worker", "6", "16 vCPU", "32 GB", "Orchestration; Spark job submit"),
    ("Spark Driver", "2-4", "8 vCPU", "32 GB", "Kubernetes cluster mode"),
    ("Spark Compute", "6-10", "48 vCPU", "256 GB", "3.84 TB NVMe local shuffle; dynamic allocation"),
    ("MinIO", "8", "24-32 vCPU", "128 GB", "8 x 12 TB enterprise disk/node; XFS, EC"),
    ("ClickHouse", "6", "48 vCPU", "256 GB", "7.68-15.36 TB NVMe/node; 3 shard x 2 replica"),
    ("ClickHouse Keeper", "3", "8 vCPU", "16 GB", "Alohida disk/node; quorum"),
    ("Trino Coordinator", "1 + standby", "16 vCPU", "64 GB", "Faqat planning va resource groups"),
    ("Trino Worker", "6", "32 vCPU", "256 GB", "2 TB NVMe spill/cache; katta executor"),
    ("Iceberg REST/JDBC Catalog", "3", "8 vCPU", "32 GB", "Table metadata, snapshot va schema evolution"),
    ("Superset Web", "4", "8 vCPU", "16 GB", "Gunicorn async, shared metadata DB"),
    ("Superset Celery", "6", "8 vCPU", "16 GB", "Async query, report va cache warmup"),
    ("Prometheus/Grafana", "3", "16 vCPU", "64 GB", "Metric va alert; 2 TB/node"),
    ("OpenSearch/ELK", "5", "16 vCPU", "64 GB", "Log va audit; 4 TB/node"),
]
add_table(
    document,
    ["Servis", "Replica", "CPU", "RAM", "Disk / vazifa"],
    service_rows,
    [Inches(1.52), Inches(0.63), Inches(0.75), Inches(0.78), Inches(2.58)],
)

document.add_heading("5. Fizik yoki virtual serverlarni guruhlash", level=1)
add_table(
    document,
    ["Server guruhi", "Soni", "Har bir server", "Workload"],
    [
        ("Control plane", "3", "16 vCPU, 64 GB, 500 GB SSD", "Kubernetes/OpenShift boshqaruvi"),
        ("Application pool", "4", "32 vCPU, 128 GB, 1.92 TB NVMe", "Next.js, FastAPI, Superset, Keycloak, Redis"),
        ("ClickHouse pool", "6", "48 vCPU, 256 GB, 7.68-15.36 TB NVMe", "3 shard x 2 replica DWH"),
        ("Shared compute pool", "8", "48 vCPU, 256 GB, 3.84 TB NVMe", "Spark, Trino, Airflow; quota bilan"),
        ("Streaming/infra pool", "4", "32 vCPU, 128 GB, 4 TB NVMe", "Kafka, NiFi, PostgreSQL, monitoring"),
        ("MinIO storage pool", "8", "24-32 vCPU, 128 GB, 8 x 12 TB", "768 TB raw object storage"),
    ],
    [Inches(1.38), Inches(0.48), Inches(2.2), Inches(2.65)],
)
document.add_paragraph(
    "Jami boshlang'ich topologiya: 3 control-plane + 22 Kubernetes worker + 8 MinIO storage node. "
    "VM ishlatilsa ham ClickHouse, Kafka va MinIO disklariga dedicated IOPS va network bandwidth kafolatlanishi kerak. "
    "Spark va Trino bir compute poolda ishlasa, Kubernetes quota va vaqt bo'yicha scheduling bilan resurs talashuvi cheklanadi."
)

document.add_heading("6. Katta data uchun to'g'ri ishlov berish modeli", level=1)
for item in [
    "FastAPI katta faylni RAMga olmaydi; client presigned multipart URL orqali MinIO'ga bevosita yuklaydi.",
    "API 202 Accepted va job_id qaytaradi; status WebSocket yoki polling orqali olinadi.",
    "NiFi oqimni route qiladi va Kafka'ga faqat event/metadata beradi; yuzlab GB payload Kafka ichiga solinmaydi.",
    "Spark fayllarni partition bo'yicha o'qiydi; collect(), toPandas() va bitta process RAMiga to'liq yig'ish taqiqlanadi.",
    "Raw, Prepared va Curated qatlamlar Parquet/Iceberg table sifatida saqlanadi; schema va partition evolyutsiyasi metadata orqali boshqariladi.",
    "Kichik fayllar compaction qilinadi; maqsadli Parquet file hajmi odatda 256-1024 MB diapazonda load-test bilan tanlanadi.",
    "ClickHouse'ga faqat dashboard va KPI uchun kerakli curated ustunlar yuklanadi; barcha raw data ko'chirilmaydi.",
    "Dashboardlar materialized view, projection va Redis/Superset cache orqali xizmat qiladi.",
    "Trino ad-hoc querylar resource group orqali navbatlanadi; bitta foydalanuvchi barcha workerlarni egallamaydi.",
    "Har bir dataset version snapshot ID, source ID, run ID, checksum va lineage bilan audit qilinadi.",
]:
    add_number(document, item)

document.add_heading("7. Data zonalari va version boshqaruvi", level=1)
add_table(
    document,
    ["Zona", "Yozuvchi servis", "O'qiydigan servis", "Version siyosati"],
    [
        ("Landing", "NiFi/FastAPI -> MinIO", "Spark validator", "Immutable original; qisqa retention"),
        ("Raw", "Spark normalize", "Spark/Trino", "Append-only; source va ingest_date partition"),
        ("Prepared", "Spark imputation/edit", "Quality/Trino", "Iceberg snapshot; manual edit yangi version"),
        ("Curated", "Spark/dbt", "ClickHouse loader/Trino", "Conformed snapshot va business version"),
        ("DWH", "Batch/stream loader", "Superset/API", "ReplicatedMergeTree; run_id va data_version"),
    ],
    [Inches(1.05), Inches(1.65), Inches(1.65), Inches(2.35)],
)

document.add_heading("8. High Availability, backup va xavfsizlik", level=1)
for item in [
    "Tashqi trafik faqat WAF/Ingress orqali; data servislar internetga ochilmaydi.",
    "Barcha ulanishlarda TLS/mTLS; secretlar Vault yoki enterprise secret manager orqali.",
    "Kafka: 5 broker, replication.factor=3, min.insync.replicas=2, acks=all.",
    "ClickHouse: 3 shard x 2 replica, 3 Keeper, alohida snapshot va object-storage backup.",
    "MinIO: 8 node, bir xil disklar, erasure coding, bucket versioning va capacity alert.",
    "PostgreSQL: 3 node HA, WAL archive, PITR, PgBouncer va alohida backup.",
    "Asosiy sayt uchun RPO 15 daqiqa va RTO 60 daqiqa target; ikkinchi sayt talabi biznes bilan tasdiqlanadi.",
    "Backup restore har chorakda test qilinadi; backup mavjudligi restore muvaffaqiyatisiz qabul qilinmaydi.",
    "Keycloak SSO/MFA/RBAC, column/row access, audit log va PII masking qo'llanadi.",
]:
    add_bullet(document, item)

document.add_heading("9. Autoscaling va back-pressure", level=1)
add_table(
    document,
    ["Workload", "Minimum", "Maximum", "Scale signali"],
    [
        ("Next.js", "4 pod", "12 pod", "RPS, CPU 60%, p95"),
        ("FastAPI", "6 pod", "24 pod", "Pending request va p95"),
        ("Job worker", "8 pod", "40 pod", "Queue depth va wait time"),
        ("Spark executor", "6 node ekvivalenti", "16 node ekvivalenti", "Pending task va shuffle"),
        ("Trino worker", "6 node", "12 node", "Queued query, CPU va memory"),
        ("Superset Web", "4 pod", "12 pod", "Active request va latency"),
        ("Superset Celery", "6 pod", "24 pod", "Celery queue depth"),
    ],
    [Inches(1.55), Inches(1.25), Inches(1.25), Inches(2.55)],
)
document.add_paragraph(
    "Storage servislar CPU asosida avtomatik pod ko'paytirish bilan emas, oldindan capacity reja va yangi server pool qo'shish orqali kengayadi. "
    "MinIO 60% fill yoki ikki yillik forecast 70% chegaradan oshishini ko'rsatsa, yangi pool xarid qilinadi."
)

document.add_heading("10. Load-test va qabul mezonlari", level=1)
for item in [
    "500 active user bilan 4 soatlik steady-state test; 800 user peak va 1000 user stress-test.",
    "Bir vaqtning o'zida dashboard, API, 8-20 Spark job va 50-100 Trino/ClickHouse query aralash workload.",
    "250 GB/kun ekvivalent ingest throughput va peak 500 GB/kun burst simulyatsiyasi.",
    "API p95 < 1 soniya, cached dashboard p95 < 3 soniya, uncached query p95 workload SLA bo'yicha.",
    "Error < 1%; CPU doimiy < 70%; RAM < 80%; MinIO fill < 70%.",
    "Bitta ClickHouse replica, Kafka broker, app node va Spark executor o'chirilganda tizim ishlashda davom etishi.",
    "500 GB va 1 TB hajmdagi fayl uchun upload, retry, resume, processing va version testlari.",
    "Backupdan to'liq restore va record count/checksum bilan ma'lumot yaxlitligi tekshirilishi.",
]:
    add_bullet(document, item)

document.add_heading("11. Joriy demodan productionga o'tish", level=1)
add_table(
    document,
    ["Joriy demo", "Katta-data production talabi"],
    [
        ("Bitta docker-compose host", "Multi-node Kubernetes/OpenShift va dedicated data poollar"),
        ("Payload Python list ichida", "Streaming/partition processing; Spark executorlar"),
        ("MinIO 1 node/drive", "8 node, 768 TB raw boshlang'ich pool"),
        ("ClickHouse 1 node", "6 node, 3 shard x 2 replica va Keeper"),
        ("Kafka 1 broker, RF=1", "5 broker, RF=3, min ISR=2"),
        ("Trino bitta process", "Coordinator + 6 katta worker"),
        ("Spark faqat kod namunasi", "Kubernetes cluster mode va dynamic allocation"),
        ("Parquet fayllar katalogsiz", "Iceberg catalog, snapshot, schema/partition evolution"),
        ("Pipeline HTTP request ichida", "Queue, job_id, worker va idempotent retry"),
        ("Superset development server", "Gunicorn, Redis cache va Celery"),
    ],
    [Inches(3.15), Inches(3.45)],
)

document.add_heading("12. Bosqichma-bosqich joriy qilish", level=1)
for item in [
    "1-bosqich: presigned multipart upload, Redis queue, job status API va idempotency.",
    "2-bosqich: Kubernetes/OpenShift, Ingress/WAF, Keycloak va observability.",
    "3-bosqich: Kafka 5 broker, NiFi cluster, PostgreSQL HA va MinIO 8 node.",
    "4-bosqich: Spark cluster, Iceberg catalog, Parquet partition va compaction.",
    "5-bosqich: ClickHouse 3 shard x 2 replica, Keeper, materialized view va backup.",
    "6-bosqich: Trino worker cluster, resource groups, Superset cache/Celery.",
    "7-bosqich: katta-data load-test, failover, restore va security acceptance.",
]:
    add_number(document, item)

document.add_heading("13. Yakuniy tavsiya", level=1)
document.add_paragraph(
    "Bu loyiha katta ma'lumot qayta ishlaydigan enterprise platforma bo'lsa, 500 ta online user uning faqat web yukidir. "
    "Asosiy sizingni kunlik ingest va retention belgilaydi. 100-250 GB/kun uchun Large-A: 22 Kubernetes worker, "
    "8 MinIO storage node va 6 ClickHouse node bilan boshlash asosli. Real ingest 250 GB/kun, safe storage 60% yoki "
    "compute queue SLA chegarasidan oshsa, yangi MinIO pool, ClickHouse shard va Spark/Trino workerlar qo'shiladi. "
    "0.5-2 TB/kun darajasida Large-B PB-scale arxitektura alohida texnik-iqtisodiy hisob bilan tasdiqlanishi kerak."
)

document.add_heading("14. Rasmiy texnik manbalar", level=1)
sources = [
    "MinIO capacity planning: https://min.io/docs/minio/linux/operations/install-deploy-manage/expand-minio-deployment.html",
    "MinIO distributed deployment: https://min.io/docs/minio/linux/operations/install-deploy-manage/deploy-minio-multi-node-multi-drive.html",
    "Apache Spark cluster overview: https://spark.apache.org/docs/latest/cluster-overview.html",
    "Apache Spark job scheduling: https://spark.apache.org/docs/latest/job-scheduling",
    "Apache Iceberg: https://iceberg.apache.org/docs/latest/",
    "Apache Iceberg partitioning: https://iceberg.apache.org/docs/latest/partitioning/",
    "Trino Iceberg connector: https://trino.io/docs/current/connector/iceberg.html",
    "Trino deployment: https://trino.io/docs/current/installation/deployment.html",
    "Apache Superset production configuration: https://superset.apache.org/admin-docs/configuration/configuring-superset/",
    "Apache Kafka broker configuration: https://kafka.apache.org/37/configuration/broker-configs/",
    "Kubernetes autoscaling: https://kubernetes.io/docs/concepts/workloads/autoscaling/",
]
for source in sources:
    add_bullet(document, source)

document.core_properties.title = "500 ta online foydalanuvchi va katta data uchun server talablari"
document.core_properties.subject = "Enterprise Data Warehouse va Data Lake production sizing"
document.core_properties.author = "Data Warehouse loyiha jamoasi"
document.core_properties.keywords = "Big Data, Data Warehouse, 500 online users, MinIO, Spark, ClickHouse, Trino, Iceberg"

document.save(OUTPUT)
print(OUTPUT)